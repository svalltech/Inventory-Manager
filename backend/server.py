from fastapi import FastAPI, APIRouter, HTTPException, Depends, status
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict, EmailStr
from typing import List, Optional, Dict
import uuid
from datetime import datetime, timezone, timedelta
import bcrypt
import jwt
from enum import Enum
from io import BytesIO
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from fastapi.responses import StreamingResponse
from fastapi import UploadFile, File
from openpyxl import load_workbook


ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# JWT Configuration
JWT_SECRET = os.environ.get('JWT_SECRET', 'your-secret-key-change-in-production')
JWT_ALGORITHM = 'HS256'
JWT_EXPIRATION_HOURS = 24

# Create the main app without a prefix
app = FastAPI(title="Inventory Management API", version="1.0.0")

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")

security = HTTPBearer()

# Enums
class UserRole(str, Enum):
    ADMIN = "admin"
    STAFF = "staff"
    VIEWER = "viewer"

class Gender(str, Enum):
    MALE = "male"
    FEMALE = "female"

class ItemStatus(str, Enum):
    ACTIVE = "active"
    DISCONTINUED = "discontinued"

class SyncStatus(str, Enum):
    SYNCED = "synced"
    PENDING_SYNC = "pending_sync"
    CONFLICT = "conflict"

# Models
class User(BaseModel):
    model_config = ConfigDict(extra="ignore")
    
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    email: EmailStr
    role: UserRole
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class UserCreate(BaseModel):
    email: EmailStr
    password: str
    role: UserRole = UserRole.VIEWER

class UserLogin(BaseModel):
    email: EmailStr
    password: str

class Token(BaseModel):
    access_token: str
    token_type: str
    user: User

class FabricSpecs(BaseModel):
    material: str
    weight: Optional[str] = None
    composition: Optional[str] = None

class InventoryItem(BaseModel):
    model_config = ConfigDict(extra="ignore")
    
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    sku: str
    name: str
    brand: str
    warehouse: str
    category: str
    gender: Gender
    color: str
    color_code: Optional[str] = None
    fabric_specs: FabricSpecs
    size: str
    design: str
    mrp: float
    selling_price: float
    cost_price: Optional[float] = None
    quantity: int
    low_stock_threshold: int = 10
    images: List[str] = []
    status: ItemStatus = ItemStatus.ACTIVE
    sync_status: SyncStatus = SyncStatus.SYNCED
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    updated_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    created_by: str
    last_modified_by: str
    last_synced_at: Optional[datetime] = None

class InventoryItemCreate(BaseModel):
    sku: str
    name: str
    brand: str
    warehouse: str
    category: str
    gender: Gender
    color: str
    color_code: Optional[str] = None
    fabric_specs: FabricSpecs
    size: str
    design: str
    mrp: float
    selling_price: float
    cost_price: Optional[float] = None
    quantity: int
    low_stock_threshold: int = 10
    images: List[str] = []
    status: ItemStatus = ItemStatus.ACTIVE

class InventoryItemUpdate(BaseModel):
    sku: Optional[str] = None
    name: Optional[str] = None
    brand: Optional[str] = None
    warehouse: Optional[str] = None
    category: Optional[str] = None
    gender: Optional[Gender] = None
    color: Optional[str] = None
    color_code: Optional[str] = None
    fabric_specs: Optional[FabricSpecs] = None
    size: Optional[str] = None
    design: Optional[str] = None
    mrp: Optional[float] = None
    selling_price: Optional[float] = None
    cost_price: Optional[float] = None
    quantity: Optional[int] = None
    low_stock_threshold: Optional[int] = None
    images: Optional[List[str]] = None
    status: Optional[ItemStatus] = None

class InventoryStats(BaseModel):
    total_items: int
    total_quantity: int
    low_stock_items: int
    categories_count: int
    total_value: float

class ExportTemplate(BaseModel):
    model_config = ConfigDict(extra="ignore")
    
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str
    fields: List[str]
    is_default: bool = False
    created_by: str
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class ExportTemplateCreate(BaseModel):
    name: str
    fields: List[str]
    is_default: bool = False

class ExportRequest(BaseModel):
    format: str  # excel, pdf, word
    fields: List[str]
    filters: Optional[Dict] = None

class ImportResult(BaseModel):
    total_rows: int
    successful: int
    failed: int
    errors: List[Dict]
    inserted: int
    updated: int

# Helper Functions
def hash_password(password: str) -> str:
    return bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')

def verify_password(password: str, hashed: str) -> bool:
    return bcrypt.checkpw(password.encode('utf-8'), hashed.encode('utf-8'))

def create_access_token(user_id: str, email: str, role: str) -> str:
    expiration = datetime.now(timezone.utc) + timedelta(hours=JWT_EXPIRATION_HOURS)
    payload = {
        'user_id': user_id,
        'email': email,
        'role': role,
        'exp': expiration
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALGORITHM)

async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)) -> dict:
    try:
        token = credentials.credentials
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        return payload
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token has expired")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")

def require_role(allowed_roles: List[UserRole]):
    async def role_checker(current_user: dict = Depends(get_current_user)):
        if current_user['role'] not in [role.value for role in allowed_roles]:
            raise HTTPException(
                status_code=403,
                detail=f"Access denied. Required roles: {[role.value for role in allowed_roles]}"
            )
        return current_user
    return role_checker

# Auth Routes
@api_router.post("/auth/register", response_model=User)
async def register(user_data: UserCreate):
    # Check if user already exists
    existing_user = await db.users.find_one({"email": user_data.email})
    if existing_user:
        raise HTTPException(status_code=400, detail="Email already registered")
    
    # Create user
    user_dict = {
        "id": str(uuid.uuid4()),
        "email": user_data.email,
        "password_hash": hash_password(user_data.password),
        "role": user_data.role.value,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.users.insert_one(user_dict)
    
    return User(
        id=user_dict["id"],
        email=user_dict["email"],
        role=user_dict["role"],
        created_at=datetime.fromisoformat(user_dict["created_at"])
    )

@api_router.post("/auth/login", response_model=Token)
async def login(credentials: UserLogin):
    # Find user
    user_doc = await db.users.find_one({"email": credentials.email})
    if not user_doc:
        raise HTTPException(status_code=401, detail="Invalid email or password")
    
    # Verify password
    if not verify_password(credentials.password, user_doc["password_hash"]):
        raise HTTPException(status_code=401, detail="Invalid email or password")
    
    # Create token
    access_token = create_access_token(user_doc["id"], user_doc["email"], user_doc["role"])
    
    user = User(
        id=user_doc["id"],
        email=user_doc["email"],
        role=user_doc["role"],
        created_at=datetime.fromisoformat(user_doc["created_at"])
    )
    
    return Token(
        access_token=access_token,
        token_type="bearer",
        user=user
    )

@api_router.get("/auth/me", response_model=User)
async def get_me(current_user: dict = Depends(get_current_user)):
    user_doc = await db.users.find_one({"id": current_user["user_id"]})
    if not user_doc:
        raise HTTPException(status_code=404, detail="User not found")
    
    return User(
        id=user_doc["id"],
        email=user_doc["email"],
        role=user_doc["role"],
        created_at=datetime.fromisoformat(user_doc["created_at"])
    )

# Inventory Routes
@api_router.post("/inventory", response_model=InventoryItem)
async def create_inventory_item(
    item_data: InventoryItemCreate,
    current_user: dict = Depends(require_role([UserRole.ADMIN, UserRole.STAFF]))
):
    # Check if SKU already exists
    existing_item = await db.inventory.find_one({"sku": item_data.sku})
    if existing_item:
        raise HTTPException(status_code=400, detail=f"SKU '{item_data.sku}' already exists")
    
    # Create item
    item = InventoryItem(
        **item_data.model_dump(),
        created_by=current_user["email"],
        last_modified_by=current_user["email"]
    )
    
    item_dict = item.model_dump()
    # Serialize datetime fields
    item_dict["created_at"] = item_dict["created_at"].isoformat()
    item_dict["updated_at"] = item_dict["updated_at"].isoformat()
    if item_dict.get("last_synced_at"):
        item_dict["last_synced_at"] = item_dict["last_synced_at"].isoformat()
    
    await db.inventory.insert_one(item_dict)
    
    return item

@api_router.get("/inventory", response_model=List[InventoryItem])
async def get_inventory(
    category: Optional[str] = None,
    gender: Optional[str] = None,
    color: Optional[str] = None,
    size: Optional[str] = None,
    min_price: Optional[float] = None,
    max_price: Optional[float] = None,
    status: Optional[ItemStatus] = None,
    search: Optional[str] = None,
    sku: Optional[str] = None,
    name: Optional[str] = None,
    design: Optional[str] = None,
    sort_by: Optional[str] = "created_at",
    sort_order: Optional[str] = "desc",
    current_user: dict = Depends(get_current_user)
):
    # Build query
    query = {}
    
    if category:
        query["category"] = category
    if gender:
        query["gender"] = gender
    if color:
        query["color"] = color
    if size:
        query["size"] = size
    if status:
        query["status"] = status.value
    if min_price is not None or max_price is not None:
        price_query = {}
        if min_price is not None:
            price_query["$gte"] = min_price
        if max_price is not None:
            price_query["$lte"] = max_price
        query["$or"] = [
            {"mrp": price_query},
            {"price": price_query}
        ]
    
    # Specific field search
    if sku:
        query["sku"] = {"$regex": sku, "$options": "i"}
    elif name:
        query["name"] = {"$regex": name, "$options": "i"}
    elif design:
        query["design"] = {"$regex": design, "$options": "i"}
    elif search:
        query["$or"] = [
            {"sku": {"$regex": search, "$options": "i"}},
            {"name": {"$regex": search, "$options": "i"}},
            {"design": {"$regex": search, "$options": "i"}}
        ]
    
    # Sort order
    sort_direction = -1 if sort_order == "desc" else 1
    
    # Fetch items
    items = await db.inventory.find(query, {"_id": 0}).sort(sort_by, sort_direction).to_list(1000)
    
    # Convert datetime strings back
    for item in items:
        if isinstance(item["created_at"], str):
            item["created_at"] = datetime.fromisoformat(item["created_at"])
        if isinstance(item["updated_at"], str):
            item["updated_at"] = datetime.fromisoformat(item["updated_at"])
        if item.get("last_synced_at") and isinstance(item["last_synced_at"], str):
            item["last_synced_at"] = datetime.fromisoformat(item["last_synced_at"])
    
    return items

@api_router.get("/inventory/filter-options")
async def get_filter_options(
    current_user: dict = Depends(get_current_user)
):
    """Get all unique values for filter dropdowns"""
    # Get all items
    all_items = await db.inventory.find({}, {"_id": 0, "brand": 1, "warehouse": 1, "category": 1, "gender": 1, "color": 1, "size": 1, "design": 1, "fabric_specs": 1}).to_list(10000)
    
    brands = sorted(list(set(item.get("brand") for item in all_items if item.get("brand"))))
    warehouses = sorted(list(set(item.get("warehouse") for item in all_items if item.get("warehouse"))))
    categories = sorted(list(set(item.get("category") for item in all_items if item.get("category"))))
    genders = sorted(list(set(item.get("gender") for item in all_items if item.get("gender"))))
    colors = sorted(list(set(item.get("color") for item in all_items if item.get("color"))))
    sizes = sorted(list(set(item.get("size") for item in all_items if item.get("size"))))
    designs = sorted(list(set(item.get("design") for item in all_items if item.get("design"))))
    
    # Extract materials and weights from fabric_specs
    materials = []
    weights = []
    for item in all_items:
        fabric = item.get("fabric_specs", {})
        if fabric.get("material"):
            materials.append(fabric["material"])
        if fabric.get("weight"):
            weights.append(fabric["weight"])
    
    materials = sorted(list(set(materials)))
    weights = sorted(list(set(weights)))
    
    return {
        "brands": brands,
        "warehouses": warehouses,
        "categories": categories,
        "genders": genders,
        "colors": colors,
        "sizes": sizes,
        "designs": designs,
        "materials": materials,
        "weights": weights
    }

@api_router.get("/inventory/brand-warehouses")
async def get_brand_warehouses(
    current_user: dict = Depends(get_current_user)
):
    """Get all brands with their associated warehouses"""
    # Get all items
    all_items = await db.inventory.find({}, {"_id": 0, "brand": 1, "warehouse": 1}).to_list(10000)
    
    # Build brand-warehouse mapping
    brand_warehouses = {}
    for item in all_items:
        brand = item.get("brand")
        warehouse = item.get("warehouse")
        if brand and warehouse:
            if brand not in brand_warehouses:
                brand_warehouses[brand] = set()
            brand_warehouses[brand].add(warehouse)
    
    # Convert sets to sorted lists
    result = {brand: sorted(list(warehouses)) for brand, warehouses in brand_warehouses.items()}
    
    return result

@api_router.get("/inventory/download-template")
async def download_import_template(
    current_user: dict = Depends(get_current_user)
):
    """Download Excel template for bulk import"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Inventory Import Template"
    
    # Headers
    headers = [
        "SKU", "Name", "Brand", "Warehouse", "Category", "Gender", "Color", "Color Code",
        "Size", "Design", "Material", "Weight", "Composition", "MRP", "Selling Price",
        "Cost Price", "Quantity", "Low Stock Threshold", "Status"
    ]
    
    # Header style
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=12)
    
    for col_idx, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_idx, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center")
    
    # Add sample data
    sample_data = [
        "SAMPLE-001", "Sample T-Shirt", "Nike", "Main Warehouse", "T-Shirt", "male",
        "Blue", "#0000FF", "M(40)", "Solid", "100% Cotton", "180", "Cotton 100%",
        1299, 999, 650, 100, 10, "active"
    ]
    
    for col_idx, value in enumerate(sample_data, 1):
        ws.cell(row=2, column=col_idx, value=value)
    
    # Auto-adjust column widths
    for col in ws.columns:
        max_length = 0
        column = col[0].column_letter
        for cell in col:
            if cell.value:
                max_length = max(max_length, len(str(cell.value)))
        ws.column_dimensions[column].width = min(max_length + 2, 30)
    
    # Save to BytesIO
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=inventory_import_template.xlsx"}
    )

@api_router.get("/inventory/{item_id}", response_model=InventoryItem)
async def get_inventory_item(
    item_id: str,
    current_user: dict = Depends(get_current_user)
):
    item = await db.inventory.find_one({"id": item_id}, {"_id": 0})
    if not item:
        raise HTTPException(status_code=404, detail="Item not found")
    
    # Convert datetime strings
    if isinstance(item["created_at"], str):
        item["created_at"] = datetime.fromisoformat(item["created_at"])
    if isinstance(item["updated_at"], str):
        item["updated_at"] = datetime.fromisoformat(item["updated_at"])
    if item.get("last_synced_at") and isinstance(item["last_synced_at"], str):
        item["last_synced_at"] = datetime.fromisoformat(item["last_synced_at"])
    
    return item

@api_router.put("/inventory/{item_id}", response_model=InventoryItem)
async def update_inventory_item(
    item_id: str,
    item_data: InventoryItemUpdate,
    current_user: dict = Depends(require_role([UserRole.ADMIN, UserRole.STAFF]))
):
    # Check if item exists
    existing_item = await db.inventory.find_one({"id": item_id})
    if not existing_item:
        raise HTTPException(status_code=404, detail="Item not found")
    
    # Prepare update data
    update_data = {k: v for k, v in item_data.model_dump(exclude_unset=True).items() if v is not None}
    
    if update_data:
        update_data["updated_at"] = datetime.now(timezone.utc).isoformat()
        update_data["last_modified_by"] = current_user["email"]
        update_data["sync_status"] = SyncStatus.PENDING_SYNC.value
        
        # Handle nested fabric_specs - convert to dict if it's a Pydantic model
        if "fabric_specs" in update_data:
            if hasattr(update_data["fabric_specs"], "model_dump"):
                update_data["fabric_specs"] = update_data["fabric_specs"].model_dump()
            # If it's already a dict, keep it as is
        
        await db.inventory.update_one(
            {"id": item_id},
            {"$set": update_data}
        )
    
    # Fetch and return updated item
    updated_item = await db.inventory.find_one({"id": item_id}, {"_id": 0})
    
    # Convert datetime strings
    if isinstance(updated_item["created_at"], str):
        updated_item["created_at"] = datetime.fromisoformat(updated_item["created_at"])
    if isinstance(updated_item["updated_at"], str):
        updated_item["updated_at"] = datetime.fromisoformat(updated_item["updated_at"])
    if updated_item.get("last_synced_at") and isinstance(updated_item["last_synced_at"], str):
        updated_item["last_synced_at"] = datetime.fromisoformat(updated_item["last_synced_at"])
    
    return updated_item

@api_router.delete("/inventory/{item_id}")
async def delete_inventory_item(
    item_id: str,
    current_user: dict = Depends(require_role([UserRole.ADMIN]))
):
    result = await db.inventory.delete_one({"id": item_id})
    
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Item not found")
    
    return {"message": "Item deleted successfully", "id": item_id}

@api_router.get("/inventory/stats/summary", response_model=InventoryStats)
async def get_inventory_stats(
    current_user: dict = Depends(get_current_user)
):
    # Get all items
    all_items = await db.inventory.find({}, {"_id": 0}).to_list(10000)
    
    total_items = len(all_items)
    total_quantity = sum(item["quantity"] for item in all_items)
    low_stock_items = len([item for item in all_items if item["quantity"] <= item.get("low_stock_threshold", 10)])
    categories = set(item["category"] for item in all_items)
    categories_count = len(categories)
    total_value = sum(item.get("selling_price", item.get("price", 0)) * item["quantity"] for item in all_items)
    
    return InventoryStats(
        total_items=total_items,
        total_quantity=total_quantity,
        low_stock_items=low_stock_items,
        categories_count=categories_count,
        total_value=round(total_value, 2)
    )

# Export Templates
@api_router.post("/export-templates", response_model=ExportTemplate)
async def create_export_template(
    template_data: ExportTemplateCreate,
    current_user: dict = Depends(get_current_user)
):
    """Create a new export template"""
    template = ExportTemplate(
        **template_data.model_dump(),
        created_by=current_user["email"]
    )
    
    template_dict = template.model_dump()
    template_dict["created_at"] = template_dict["created_at"].isoformat()
    
    await db.export_templates.insert_one(template_dict)
    
    return template

@api_router.get("/export-templates", response_model=List[ExportTemplate])
async def get_export_templates(
    current_user: dict = Depends(get_current_user)
):
    """Get all export templates for current user"""
    templates = await db.export_templates.find(
        {"created_by": current_user["email"]},
        {"_id": 0}
    ).to_list(100)
    
    for template in templates:
        if isinstance(template["created_at"], str):
            template["created_at"] = datetime.fromisoformat(template["created_at"])
    
    return templates

@api_router.delete("/export-templates/{template_id}")
async def delete_export_template(
    template_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Delete an export template"""
    result = await db.export_templates.delete_one({
        "id": template_id,
        "created_by": current_user["email"]
    })
    
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Template not found")
    
    return {"message": "Template deleted successfully"}

# Export Functions
def generate_excel(items: List[dict], fields: List[str]) -> BytesIO:
    """Generate Excel file"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Inventory Report"
    
    # Header style
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=12)
    
    # Write headers
    for col_idx, field in enumerate(fields, 1):
        cell = ws.cell(row=1, column=col_idx, value=field.replace("_", " ").title())
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")
    
    # Write data
    for row_idx, item in enumerate(items, 2):
        for col_idx, field in enumerate(fields, 1):
            value = item.get(field, "")
            if field == "fabric_specs":
                value = item.get("fabric_specs", {}).get("material", "")
            ws.cell(row=row_idx, column=col_idx, value=str(value))
    
    # Auto-adjust column widths
    for col in ws.columns:
        max_length = 0
        column = col[0].column_letter
        for cell in col:
            if cell.value:
                max_length = max(max_length, len(str(cell.value)))
        ws.column_dimensions[column].width = min(max_length + 2, 50)
    
    # Save to BytesIO
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    return output

def generate_word(items: List[dict], fields: List[str]) -> BytesIO:
    """Generate Word document"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Inventory Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata
    doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(f'Total Items: {len(items)}')
    doc.add_paragraph('')
    
    # Create table
    table = doc.add_table(rows=1, cols=len(fields))
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    for idx, field in enumerate(fields):
        header_cells[idx].text = field.replace("_", " ").title()
        # Bold header
        for paragraph in header_cells[idx].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Data rows
    for item in items:
        row_cells = table.add_row().cells
        for idx, field in enumerate(fields):
            value = item.get(field, "")
            if field == "fabric_specs":
                value = item.get("fabric_specs", {}).get("material", "")
            row_cells[idx].text = str(value)
    
    # Save to BytesIO
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def generate_pdf(items: List[dict], fields: List[str]) -> BytesIO:
    """Generate PDF document"""
    output = BytesIO()
    doc = SimpleDocTemplate(output, pagesize=A4)
    elements = []
    
    # Styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#4472C4'),
        spaceAfter=30,
        alignment=1  # Center
    )
    
    # Title
    elements.append(Paragraph("Inventory Report", title_style))
    elements.append(Spacer(1, 0.2*inch))
    
    # Metadata
    meta_style = styles['Normal']
    elements.append(Paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", meta_style))
    elements.append(Paragraph(f"Total Items: {len(items)}", meta_style))
    elements.append(Spacer(1, 0.3*inch))
    
    # Prepare table data
    table_data = []
    
    # Headers
    headers = [field.replace("_", " ").title() for field in fields]
    table_data.append(headers)
    
    # Data rows
    for item in items:
        row = []
        for field in fields:
            value = item.get(field, "")
            if field == "fabric_specs":
                value = item.get("fabric_specs", {}).get("material", "")
            row.append(str(value))
        table_data.append(row)
    
    # Create table
    table = Table(table_data)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4472C4')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('FONTSIZE', (0, 1), (-1, -1), 8),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
    ]))
    
    elements.append(table)
    
    # Build PDF
    doc.build(elements)
    output.seek(0)
    return output

@api_router.post("/inventory/import", response_model=ImportResult)
async def import_inventory(
    file: UploadFile = File(...),
    current_user: dict = Depends(require_role([UserRole.ADMIN, UserRole.STAFF]))
):
    """Import inventory data from Excel file"""
    if not file.filename.endswith(('.xlsx', '.xls')):
        raise HTTPException(status_code=400, detail="Only Excel files (.xlsx, .xls) are supported")
    
    try:
        # Read Excel file
        contents = await file.read()
        wb = load_workbook(BytesIO(contents))
        ws = wb.active
        
        # Get headers from first row
        headers = []
        for cell in ws[1]:
            if cell.value:
                headers.append(str(cell.value).lower().replace(" ", "_"))
        
        # Required fields
        required_fields = ['sku', 'name', 'brand', 'warehouse', 'category', 'gender', 'size', 'design', 'mrp', 'selling_price', 'quantity']
        missing_fields = [field for field in required_fields if field not in headers]
        
        if missing_fields:
            raise HTTPException(
                status_code=400,
                detail=f"Missing required columns: {', '.join(missing_fields)}"
            )
        
        # Process rows
        total_rows = 0
        successful = 0
        failed = 0
        inserted = 0
        updated = 0
        errors = []
        
        for row_idx, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
            if not any(row):  # Skip empty rows
                continue
            
            total_rows += 1
            
            try:
                # Build item dict from row
                item_data = {}
                for col_idx, value in enumerate(row):
                    if col_idx < len(headers):
                        header = headers[col_idx]
                        item_data[header] = value
                
                # Validate required fields
                for field in required_fields:
                    if not item_data.get(field):
                        raise ValueError(f"Missing required field: {field}")
                
                # Handle fabric_specs
                fabric_specs = {
                    "material": str(item_data.get("material", "")),
                    "weight": str(item_data.get("weight", "")) if item_data.get("weight") else None,
                    "composition": str(item_data.get("composition", "")) if item_data.get("composition") else None
                }
                
                # Check if item exists (by SKU)
                existing_item = await db.inventory.find_one({"sku": item_data["sku"]})
                
                if existing_item:
                    # Update existing item
                    update_data = {
                        "name": str(item_data["name"]),
                        "brand": str(item_data["brand"]),
                        "warehouse": str(item_data["warehouse"]),
                        "category": str(item_data["category"]),
                        "gender": str(item_data["gender"]).lower(),
                        "color": str(item_data.get("color", "")),
                        "color_code": str(item_data.get("color_code", "")) if item_data.get("color_code") else None,
                        "fabric_specs": fabric_specs,
                        "size": str(item_data["size"]),
                        "design": str(item_data["design"]),
                        "mrp": float(item_data["mrp"]),
                        "selling_price": float(item_data["selling_price"]),
                        "cost_price": float(item_data["cost_price"]) if item_data.get("cost_price") else None,
                        "quantity": int(item_data["quantity"]),
                        "low_stock_threshold": int(item_data.get("low_stock_threshold", 10)),
                        "status": str(item_data.get("status", "active")).lower(),
                        "updated_at": datetime.now(timezone.utc).isoformat(),
                        "last_modified_by": current_user["email"]
                    }
                    
                    await db.inventory.update_one(
                        {"sku": item_data["sku"]},
                        {"$set": update_data}
                    )
                    updated += 1
                    successful += 1
                else:
                    # Insert new item
                    new_item = {
                        "id": str(uuid.uuid4()),
                        "sku": str(item_data["sku"]),
                        "name": str(item_data["name"]),
                        "brand": str(item_data["brand"]),
                        "warehouse": str(item_data["warehouse"]),
                        "category": str(item_data["category"]),
                        "gender": str(item_data["gender"]).lower(),
                        "color": str(item_data.get("color", "")),
                        "color_code": str(item_data.get("color_code", "")) if item_data.get("color_code") else None,
                        "fabric_specs": fabric_specs,
                        "size": str(item_data["size"]),
                        "design": str(item_data["design"]),
                        "mrp": float(item_data["mrp"]),
                        "selling_price": float(item_data["selling_price"]),
                        "cost_price": float(item_data["cost_price"]) if item_data.get("cost_price") else None,
                        "quantity": int(item_data["quantity"]),
                        "low_stock_threshold": int(item_data.get("low_stock_threshold", 10)),
                        "images": [],
                        "status": str(item_data.get("status", "active")).lower(),
                        "sync_status": "synced",
                        "created_at": datetime.now(timezone.utc).isoformat(),
                        "updated_at": datetime.now(timezone.utc).isoformat(),
                        "created_by": current_user["email"],
                        "last_modified_by": current_user["email"],
                        "last_synced_at": None
                    }
                    
                    await db.inventory.insert_one(new_item)
                    inserted += 1
                    successful += 1
                
            except Exception as e:
                failed += 1
                errors.append({
                    "row": row_idx,
                    "sku": item_data.get("sku", "Unknown"),
                    "error": str(e)
                })
        
        return ImportResult(
            total_rows=total_rows,
            successful=successful,
            failed=failed,
            errors=errors,
            inserted=inserted,
            updated=updated
        )
        
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to process Excel file: {str(e)}")

@api_router.post("/inventory/export")
async def export_inventory(
    export_request: ExportRequest,
    current_user: dict = Depends(get_current_user)
):
    """Export inventory data in specified format"""
    # Get filtered items
    filters = export_request.filters or {}
    query = {}
    
    if filters.get("brand"):
        query["brand"] = filters["brand"]
    if filters.get("warehouse"):
        query["warehouse"] = filters["warehouse"]
    if filters.get("category"):
        query["category"] = filters["category"]
    if filters.get("gender"):
        query["gender"] = filters["gender"]
    
    items = await db.inventory.find(query, {"_id": 0}).to_list(10000)
    
    # Convert datetime objects to strings
    for item in items:
        if isinstance(item.get("created_at"), datetime):
            item["created_at"] = item["created_at"].strftime("%Y-%m-%d %H:%M:%S")
        if isinstance(item.get("updated_at"), datetime):
            item["updated_at"] = item["updated_at"].strftime("%Y-%m-%d %H:%M:%S")
    
    # Generate file based on format
    if export_request.format == "excel":
        file_data = generate_excel(items, export_request.fields)
        filename = f"inventory_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        media_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    elif export_request.format == "word":
        file_data = generate_word(items, export_request.fields)
        filename = f"inventory_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif export_request.format == "pdf":
        file_data = generate_pdf(items, export_request.fields)
        filename = f"inventory_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        media_type = "application/pdf"
    else:
        raise HTTPException(status_code=400, detail="Invalid format. Use 'excel', 'word', or 'pdf'")
    
    return StreamingResponse(
        file_data,
        media_type=media_type,
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

# Health check
@api_router.get("/health")
async def health_check():
    return {"status": "healthy", "service": "inventory-api", "version": "1.0.0"}

# Test UI endpoint
@api_router.get("/test-ui")
async def serve_test_ui():
    return FileResponse(str(ROOT_DIR / "static" / "test.html"))

# Include the router in the main app
app.include_router(api_router)

# Serve static files for test UI
app.mount("/static", StaticFiles(directory=str(ROOT_DIR / "static")), name="static")

@app.get("/test")
async def serve_test_page():
    return FileResponse(str(ROOT_DIR / "static" / "test.html"))

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()